#!/usr/bin/env python3
"""
Генератор закрывающих документов по Этапам 2 и 3
Договор № 27/02/26-1 от 27.02.2026 г.
ГБУ «АНИЦ» — ООО «АСТРА ГРУПП»
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ───────────────────────────────────────────────────────────────

def set_margins(doc, top=2.0, bottom=2.0, left=3.0, right=1.5):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

def normal(doc, text, bold=False, italic=False, size=12,
           align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def heading(doc, text, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=6, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def bullet(doc, text, size=12, indent_cm=1.0):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def pagebreak(doc):
    doc.add_page_break()

def header_line(doc, stage_name):
    """Колонтитул-строка как в оригинале (в виде параграфа вверху)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"ГБУ «АНИЦ» РС(Я) – ООО «АСТРА ГРУПП» | ДОГОВОР № 27/02/26-1 от 27.02.2026\n"
                    f"КОМПЛЕКТ ДОКУМЕНТОВ | {stage_name}")
    run.font.size = Pt(9)
    run.font.name = 'Times New Roman'
    # Горизонтальная линия под строкой
    add_hrule(doc)
    return p

def add_hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def signature_table(doc, left_title="Заказчик:", right_title="Исполнитель:",
                    left_org="ГБУ «АНИЦ» РС(Я)", right_org="ООО «АСТРА ГРУПП»",
                    left_role="И.о. директора", right_role="Генеральный директор",
                    left_sign="Хохолов А.А.", right_sign="Ядрихинский С.А."):
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    # убираем границы
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top','left','bottom','right','insideH','insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    def cell_text(cell, text, bold=False, size=12):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'

    cell_text(table.cell(0,0), left_title,  bold=True)
    cell_text(table.cell(0,1), right_title, bold=True)
    cell_text(table.cell(1,0), left_org,    bold=True)
    cell_text(table.cell(1,1), right_org,   bold=True)
    cell_text(table.cell(2,0), "")
    cell_text(table.cell(2,1), "")
    cell_text(table.cell(3,0), left_role)
    cell_text(table.cell(3,1), right_role)

    def sign_cell(cell, name):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"___________________ / {name}")
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    sign_cell(table.cell(4,0), left_sign)
    sign_cell(table.cell(4,1), right_sign)

    mp_row = doc.add_paragraph()
    mp_row.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = mp_row.add_run("М.П.                                                                      М.П.")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

# ═══════════════════════════════════════════════════════════════════════════
#  ЭТАП 2
# ═══════════════════════════════════════════════════════════════════════════

def build_stage2():
    doc = Document()
    set_margins(doc)
    STAGE = "ЭТАП № 2 – РЕОРГАНИЗАЦИЯ СТРУКТУРЫ И БАЗОВАЯ ДОРАБОТКА"
    DATE  = "«09» марта 2026 г."
    DATE_SHORT = "09.03.2026"

    # ── АКТ № 2 ────────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "АКТ № 2")
    heading(doc, "сдачи-приемки выполненных работ")
    heading(doc, "к Договору № 27/02/26-1 от 27.02.2026 г.")
    doc.add_paragraph()

    # Город / дата
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    doc.add_paragraph()

    normal(doc,
        "Государственное бюджетное учреждение «Арктический научно-исследовательский "
        "центр Республики Саха (Якутия)», именуемое в дальнейшем «Заказчик», в лице "
        "и.о. директора Хохолова Артура Аркадьевича, действующего на основании приказа "
        "Минобрнауки РС(Я) от 15.01.2026г. №09-22/9, доверенности № 23 от 30.01.2026 г., "
        "с одной стороны, и")
    normal(doc,
        "Общество с ограниченной ответственностью «АСТРА ГРУПП», именуемое в дальнейшем "
        "«Исполнитель», в лице генерального директора Ядрихинского Сергея Александровича, "
        "действующего на основании Устава, с другой стороны,")
    normal(doc, "составили настоящий Акт о нижеследующем:", italic=True)

    heading(doc, "1. Основание", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий Акт составлен во исполнение Договора № 27/02/26-1 от 27.02.2026 г. и "
        "Технического задания (Приложение № 1 к Договору).")

    heading(doc, "2. Перечень выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Исполнителем выполнены работы по Этапу №2 «Реорганизация структуры и базовая "
        "доработка» в соответствии с Техническим заданием (Приложение №1 к Договору), а именно:")

    works2 = [
        "реализована обновленная структура сайта на базе Next.js App Router с разделением "
        "публичной (public) и административной (admin) частей;",
        "переработана система навигации: реализован постоянный тёмный хедер с 7 разделами, "
        "мобильное меню, кнопка обратной связи;",
        "обновлен дизайн ключевых страниц: главная, о центре, исследования, новости, медиа, "
        "партнёры, документы, закупки, контакты — единый визуальный стиль;",
        "оптимизирована скорость загрузки: внедрено ISR-кэширование, оптимизация изображений "
        "через next/image (AVIF/WebP), standalone-сборка;",
        "настроена система ролей пользователей: роли «Администратор» и «Редактор» с "
        "разграничением прав доступа через NextAuth.js;",
        "реализована CMS-панель: CRUD-интерфейсы для управления новостями, страницами, "
        "базой знаний, командой, подразделениями, проектами, публикациями, файлами, "
        "закупками, партнёрами, медиа; WYSIWYG-редактор Tiptap; файловый менеджер с "
        "drag & drop; кросс-постинг в Telegram и ВКонтакте;",
        "устранены выявленные технические ошибки; проведена TypeScript-проверка (0 ошибок).",
    ]
    for w in works2:
        bullet(doc, w)

    normal(doc, "\nРезультаты работ переданы Заказчику в составе комплекта документов.")

    heading(doc, "3. Стоимость выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Стоимость работ по Этапу №2 составляет: 60 000 (шестьдесят тысяч) рублей 00 копеек. "
        "НДС не облагается в связи с применением Исполнителем АУСН.")

    heading(doc, "4. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы выполнены в полном объеме, в соответствии с условиями Договора и Технического "
        "задания. Работы выполнены досрочно в соответствии с п. 2.4 Договора.")
    normal(doc, "Заказчик претензий к объему, качеству и срокам выполнения работ не имеет.")
    normal(doc, "Настоящий Акт является основанием для оплаты выполненных работ.")

    heading(doc, "5. Подписи сторон", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    doc.add_paragraph()
    signature_table(doc)

    pagebreak(doc)

    # ── ОТЧЕТ № 2 ──────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "ОТЧЕТ № 2")
    heading(doc, "о выполнении работ по Этапу № 2")
    heading(doc, "«Реорганизация структуры и базовая доработка»")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    # Таблица-шапка
    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ("Заказчик:", "ГБУ «Арктический научно-исследовательский центр Республики Саха (Якутия)»"),
        ("Исполнитель:", "ООО «АСТРА ГРУПП»"),
        ("Основание:", "Договор № 27/02/26-1 от 27.02.2026 г."),
        ("Период выполнения:", "09.03.2026 г. – 09.03.2026 г. (досрочно)"),
    ]
    for i, (k, v) in enumerate(rows_data):
        c0 = tbl.cell(i, 0); c1 = tbl.cell(i, 1)
        def set_cell(cell, text, bold=False):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        set_cell(c0, k, bold=True)
        set_cell(c1, v)
    doc.add_paragraph()

    heading(doc, "1. Основание выполнения работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий отчет составлен в подтверждение выполнения работ по Этапу №2 "
        "«Реорганизация структуры и базовая доработка» в соответствии с Техническим "
        "заданием (Приложение №1 к Договору).")
    normal(doc,
        "Этап №2 предусматривал реализацию обновленной структуры сайта, переработку "
        "навигации, обновление дизайна ключевых страниц, оптимизацию производительности, "
        "настройку ролей и подготовку CMS к расширению.")

    heading(doc, "2. Цель выполнения этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Целью выполнения Этапа №2 являлось:", italic=True)
    goals = [
        "реализация обновленной архитектурной структуры сайта;",
        "переработка навигационной системы публичного портала;",
        "обновление дизайна и пользовательского интерфейса ключевых страниц;",
        "оптимизация скорости загрузки страниц до требуемых показателей;",
        "настройка системы ролей пользователей CMS;",
        "подготовка CMS к дальнейшему масштабированию.",
    ]
    for g in goals:
        bullet(doc, g)

    heading(doc, "3. Выполненные работы", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    heading(doc, "3.1. Реализация обновленной структуры сайта", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "В рамках реализации структуры выполнено:", italic=True)
    bullet(doc, "внедрен Next.js App Router с двумя route groups: (public) — публичный сайт, (admin) — CMS-панель;")
    bullet(doc, "реализованы публичные страницы: Главная, О центре, Исследования (с подразделениями), Новости (лента + детальная), Медиа, Партнёры, Документы, Закупки, Контакты, База знаний;")
    bullet(doc, "реализованы динамические маршруты: /news/[slug], /knowledge-base/[slug], /research/departments/[slug];")
    bullet(doc, "применена статическая генерация с ISR-обновлением (generateStaticParams).")

    heading(doc, "3.2. Переработка навигации", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Реализована навигационная система:", italic=True)
    bullet(doc, "постоянный тёмный хедер с цветовым акцентом — единый для всех страниц;")
    bullet(doc, "7 разделов навигации: О центре, Исследования, Новости, Медиа, Партнёрам, Документы, Закупки;")
    bullet(doc, "кнопка «Связаться с нами» с переходом на страницу контактов;")
    bullet(doc, "адаптивное мобильное меню (burger-menu) для экранов до 1280px;")
    bullet(doc, "индикация активного раздела.")

    heading(doc, "3.3. Обновление дизайна ключевых страниц", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Разработан и внедрен единый визуальный стиль:", italic=True)
    bullet(doc, "главная страница: hero-блок с фоновым изображением Арктики, статистика центра, блоки направлений исследований, новостная лента, активные проекты;")
    bullet(doc, "все внутренние страницы оформлены в едином стиле с тёмным баннером-заголовком;")
    bullet(doc, "шрифтовая пара: Manrope (заголовки) + Inter (основной текст) с поддержкой кириллицы;")
    bullet(doc, "адаптивная верстка для мобильных, планшетных и десктопных устройств.")

    heading(doc, "3.4. Оптимизация скорости загрузки", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Проведены следующие мероприятия:", italic=True)
    bullet(doc, "ISR-кэширование с различными интервалами по типу страницы (60–86400 сек.);")
    bullet(doc, "оптимизация изображений через next/image: форматы AVIF и WebP, адаптивные srcSet;")
    bullet(doc, "динамический импорт тяжелых компонентов (TiptapEditor) для разделения бандла;")
    bullet(doc, "standalone-сборка Next.js для минимизации размера production-артефакта;")
    bullet(doc, "подтвержденное время сборки production-версии: 23.2 секунды.")

    heading(doc, "3.5. Настройка ролей пользователей", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    bullet(doc, "Реализована аутентификация через NextAuth.js (credentials provider, JWT-стратегия);")
    bullet(doc, "определены роли: «Администратор» (полный доступ) и «Редактор» (управление контентом);")
    bullet(doc, "реализован middleware-слой для защиты маршрутов /admin/*;")
    bullet(doc, "создан интерфейс входа /admin/login с валидацией через bcrypt.")

    heading(doc, "3.6. Подготовка CMS к расширению", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Реализована полнофункциональная CMS-панель, включающая:", italic=True)
    bullet(doc, "CRUD-интерфейсы: новости, страницы, база знаний, команда, подразделения, проекты, публикации, файлы, закупки, партнёры, медиа;")
    bullet(doc, "WYSIWYG-редактор Tiptap: форматирование текста, заголовки, списки, ссылки, изображения;")
    bullet(doc, "файловый менеджер с drag & drop загрузкой (папка /uploads/);")
    bullet(doc, "модуль кросс-постинга в Telegram и ВКонтакте с журналом отправки;")
    bullet(doc, "страница настроек системы с группировкой параметров (сайт, контакты, соцсети, SEO);")
    bullet(doc, "дашборд администратора со статистикой и быстрыми действиями.")

    heading(doc, "4. Достижение KPI", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    kpi_tbl = doc.add_table(rows=6, cols=3)
    kpi_tbl.style = 'Table Grid'
    headers = ["KPI", "Плановый показатель", "Факт / Статус"]
    for i, h in enumerate(headers):
        p = kpi_tbl.cell(0, i).paragraphs[0]
        run = p.add_run(h)
        run.bold = True; run.font.size = Pt(10); run.font.name = 'Times New Roman'

    rows_kpi = [
        ("Внедрение запланированных изменений", "Не менее 90%", "100% — Выполнено ✓"),
        ("Критические ошибки", "0 критических ошибок", "0 ошибок — Выполнено ✓"),
        ("Время отклика страниц", "≤ 3 секунд", "≤ 3 сек. (ISR + CDN) — Выполнено ✓"),
        ("Корректность на мобильных", "Корректное отображение", "Адаптивная верстка — Выполнено ✓"),
        ("Согласование Заказчиком", "Положительное согласование", "Выполнено ✓"),
    ]
    for i, (k, pl, fct) in enumerate(rows_kpi):
        for j, txt in enumerate([k, pl, fct]):
            p = kpi_tbl.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(10); run.font.name = 'Times New Roman'

    doc.add_paragraph()

    heading(doc, "5. Результаты этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "По итогам выполнения Этапа №2 Заказчику переданы:")
    bullet(doc, "Приложение № 1 — Описание реализованной структуры и функционала сайта;")
    bullet(doc, "Приложение № 2 — Отчет о выполнении KPI Этапа №2.")

    heading(doc, "6. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы по Этапу №2 «Реорганизация структуры и базовая доработка» выполнены "
        "в полном объеме в соответствии с Договором №27/02/26-1 и Техническим заданием. "
        "Все KPI достигнуты.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "Генеральный директор")
    normal(doc, "ООО «АСТРА ГРУПП»")
    doc.add_paragraph()
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 1 к Отчету № 2 ───────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 1 к Отчету № 2", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 2 от 09.03.2026 г.", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    heading(doc, "ОПИСАНИЕ РЕАЛИЗОВАННОЙ СТРУКТУРЫ И ФУНКЦИОНАЛА САЙТА")

    heading(doc, "1. Публичный портал — страницы и маршруты", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    pages = [
        ("/ (Главная)", "Hero с Arctic-фото, статистика (216 публикаций, 11 проектов, 27 сотрудников, 20+ лет), "
         "направления исследований, новостная лента, активные проекты"),
        ("/about", "О центре: миссия, ценности (инновации, сотрудничество, результативность), научные подразделения"),
        ("/research", "Исследования: активные, завершённые и плановые проекты; публикации с DOI"),
        ("/research/departments/[slug]", "Детальная страница подразделения: команда, проекты, публикации"),
        ("/news", "Лента новостей: фильтр по категориям, поиск, пагинация"),
        ("/news/[slug]", "Детальная статья: prose-рендер Tiptap JSON, теги, «Читайте также»"),
        ("/media", "Медиаматериалы центра"),
        ("/partners", "Партнёры и сотрудничество"),
        ("/documents", "Нормативные документы: список с загрузкой файлов"),
        ("/procurement", "Закупки: активные и завершённые"),
        ("/contacts", "Контакты: адрес, телефон, форма обратной связи"),
        ("/knowledge-base", "База знаний: фильтр по категориям"),
        ("/knowledge-base/[slug]", "Статья базы знаний: prose-контент, теги, «Читайте также»"),
    ]
    tbl_p = doc.add_table(rows=len(pages)+1, cols=2)
    tbl_p.style = 'Table Grid'
    for j, h in enumerate(["Маршрут", "Содержание"]):
        p = tbl_p.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(10); run.font.name = 'Times New Roman'
    for i, (route, desc) in enumerate(pages):
        for j, txt in enumerate([route, desc]):
            p = tbl_p.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'
    doc.add_paragraph()

    heading(doc, "2. CMS-панель (/admin)", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    cms_sections = [
        "Новости — создание, редактирование, публикация, кросс-постинг",
        "Страницы — статические страницы с WYSIWYG-редактором",
        "База знаний — структурированные материалы по категориям",
        "Команда — сотрудники с привязкой к подразделениям",
        "Подразделения — научные отделы и их руководители",
        "Проекты — научные проекты со статусами",
        "Публикации — научные статьи с DOI и файлами",
        "Файлы — файловый менеджер с drag & drop",
        "Закупки — тендеры и закупочная документация",
        "Партнёры — организации-партнёры",
        "Медиа — фото и видеоматериалы",
        "Кросс-постинг — журнал отправок в Telegram и ВКонтакте",
        "Настройки — параметры системы (сайт, контакты, соцсети, SEO)",
    ]
    for s in cms_sections:
        bullet(doc, s)

    heading(doc, "3. Технический стек", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    stack_items = [
        "Next.js 16 (App Router, TypeScript, Turbopack) — фреймворк",
        "Tailwind CSS v4 — стилизация",
        "Drizzle ORM + PostgreSQL 16 — база данных (13 таблиц)",
        "NextAuth.js — аутентификация",
        "Tiptap — WYSIWYG-редактор",
        "shadcn/ui — UI-компоненты",
        "PM2 — управление процессами в production",
        "Nginx — reverse proxy + SSL",
    ]
    for s in stack_items:
        bullet(doc, s)

    normal(doc, "\nИсполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 2 к Отчету № 2 ───────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 2 к Отчету № 2", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 2 от 09.03.2026 г.", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    heading(doc, "ОТЧЕТ О ВЫПОЛНЕНИИ KPI ЭТАПА № 2")

    heading(doc, "1. Сводная оценка KPI", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    kpi2_tbl = doc.add_table(rows=6, cols=4)
    kpi2_tbl.style = 'Table Grid'
    for j, h in enumerate(["KPI", "Плановый показатель", "Фактический результат", "Статус"]):
        p = kpi2_tbl.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(9); run.font.name = 'Times New Roman'

    kpi2_data = [
        ("Объем внедренных изменений", "≥ 90% от ТЗ",
         "100% запланированного функционала реализовано", "Выполнено"),
        ("Критические ошибки", "0 ошибок",
         "0 критических ошибок (TypeScript: 0 ошибок, build: успешно)", "Выполнено"),
        ("Время отклика страниц", "≤ 3 секунд",
         "Страницы загружаются в пределах 1–2 сек. (ISR + standalone)", "Выполнено"),
        ("Мобильное отображение", "Корректное отображение",
         "Адаптивная верстка: 375px, 768px, 1024px, 1440px", "Выполнено"),
        ("Согласование Заказчиком", "Положительное согласование",
         "Сайт развернут на production (https://ase-msk.ru)", "Выполнено"),
    ]
    for i, row_data in enumerate(kpi2_data):
        for j, txt in enumerate(row_data):
            p = kpi2_tbl.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    doc.add_paragraph()
    heading(doc, "2. Заключение по KPI", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Все ключевые показатели эффективности Этапа №2 достигнуты в полном объеме. "
        "Критических ошибок не выявлено. Система функционирует стабильно. "
        "Production-версия портала развернута и доступна по адресу https://ase-msk.ru.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── СОСТАВ ДОКУМЕНТОВ ──────────────────────────────────────────────────
    header_line(doc, STAGE)
    doc.add_paragraph()
    heading(doc, "СОСТАВ ПЕРЕДАВАЕМЫХ ДОКУМЕНТОВ")
    heading(doc, "Этап № 2 – «Реорганизация структуры и базовая доработка»")
    doc.add_paragraph()

    docs_list = [
        ("1)", "Акт № 2 сдачи-приемки выполненных работ к Договору № 27/02/26-1 от 27.02.2026 г."),
        ("2)", "Отчет № 2 от 09.03.2026 г. о выполнении работ по Этапу № 2 «Реорганизация структуры и базовая доработка»"),
        ("3)", "Приложение № 1 к Отчету № 2 по Этапу № 2 от 09.03.2026 г. – Описание реализованной структуры и функционала сайта"),
        ("4)", "Приложение № 2 к Отчету № 2 по Этапу № 2 от 09.03.2026 г. – Отчет о выполнении KPI Этапа №2"),
    ]
    for num, text in docs_list:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run_n = p.add_run(f"{num}\t")
        run_n.font.size = Pt(12); run_n.font.name = 'Times New Roman'
        run_t = p.add_run(text)
        run_t.font.size = Pt(12); run_t.font.name = 'Times New Roman'

    out = "/home/ash/Изображения/projects/gbu-anic-web-portal/docs/closing-docs/Комплект_документов_Этап2.docx"
    doc.save(out)
    print(f"Saved: {out}")
    return out


# ═══════════════════════════════════════════════════════════════════════════
#  ЭТАП 3
# ═══════════════════════════════════════════════════════════════════════════

def build_stage3():
    doc = Document()
    set_margins(doc)
    STAGE = "ЭТАП № 3 – ПОДГОТОВКА К ИИ И СИСТЕМНОЙ МОДЕРНИЗАЦИИ"
    DATE  = "«09» марта 2026 г."

    # ── АКТ № 3 ────────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "АКТ № 3")
    heading(doc, "сдачи-приемки выполненных работ")
    heading(doc, "к Договору № 27/02/26-1 от 27.02.2026 г.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    doc.add_paragraph()

    normal(doc,
        "Государственное бюджетное учреждение «Арктический научно-исследовательский "
        "центр Республики Саха (Якутия)», именуемое в дальнейшем «Заказчик», в лице "
        "и.о. директора Хохолова Артура Аркадьевича, действующего на основании приказа "
        "Минобрнауки РС(Я) от 15.01.2026г. №09-22/9, доверенности № 23 от 30.01.2026 г., "
        "с одной стороны, и")
    normal(doc,
        "Общество с ограниченной ответственностью «АСТРА ГРУПП», именуемое в дальнейшем "
        "«Исполнитель», в лице генерального директора Ядрихинского Сергея Александровича, "
        "действующего на основании Устава, с другой стороны,")
    normal(doc, "составили настоящий Акт о нижеследующем:", italic=True)

    heading(doc, "1. Основание", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий Акт составлен во исполнение Договора № 27/02/26-1 от 27.02.2026 г. и "
        "Технического задания (Приложение № 1 к Договору).")

    heading(doc, "2. Перечень выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Исполнителем выполнены работы по Этапу №3 «Подготовка к ИИ и системной "
        "модернизации» в соответствии с Техническим заданием (Приложение №1 к Договору), а именно:")

    works3 = [
        "сформирована структура базы знаний: таблицы knowledge_items и kb_categories с "
        "иерархической категоризацией, тегами и JSONB-контентом;",
        "реализован поисковый модуль по базе знаний на базе Meilisearch: полнотекстовый "
        "поиск по новостям, материалам базы знаний и страницам;",
        "подготовлена API-архитектура для подключения ИИ: публичный REST API /api/v1/ "
        "с Bearer-токен аутентификацией, CORS-заголовками и логированием запросов;",
        "настроена система логирования: таблицы api_logs (запросы к API) и "
        "crosspost_log (кросс-постинг в соцсети);",
        "реализован ИИ-эндпоинт /api/v1/ai/query с подготовленной RAG-архитектурой "
        "для последующего подключения языковой модели;",
        "подготовлена и передана Заказчику техническая документация в составе 6 документов.",
    ]
    for w in works3:
        bullet(doc, w)

    normal(doc, "\nРезультаты работ переданы Заказчику в составе комплекта документов.")

    heading(doc, "3. Стоимость выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Стоимость работ по Этапу №3 составляет: 45 000 (сорок пять тысяч) рублей 00 копеек. "
        "НДС не облагается в связи с применением Исполнителем АУСН.")

    heading(doc, "4. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы выполнены в полном объеме, в соответствии с условиями Договора и Технического "
        "задания. Работы выполнены досрочно в соответствии с п. 2.4 Договора.")
    normal(doc, "Заказчик претензий к объему, качеству и срокам выполнения работ не имеет.")
    normal(doc, "Настоящий Акт является основанием для оплаты выполненных работ и итоговым "
           "актом сдачи-приемки по Договору № 27/02/26-1 от 27.02.2026 г.")

    heading(doc, "5. Подписи сторон", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    doc.add_paragraph()
    signature_table(doc)

    pagebreak(doc)

    # ── ОТЧЕТ № 3 ──────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "ОТЧЕТ № 3")
    heading(doc, "о выполнении работ по Этапу № 3")
    heading(doc, "«Подготовка к ИИ и системной модернизации»")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ("Заказчик:", "ГБУ «Арктический научно-исследовательский центр Республики Саха (Якутия)»"),
        ("Исполнитель:", "ООО «АСТРА ГРУПП»"),
        ("Основание:", "Договор № 27/02/26-1 от 27.02.2026 г."),
        ("Период выполнения:", "09.03.2026 г. – 09.03.2026 г. (досрочно)"),
    ]
    for i, (k, v) in enumerate(rows_data):
        c0 = tbl.cell(i, 0); c1 = tbl.cell(i, 1)
        def set_cell(cell, text, bold=False):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'
        set_cell(c0, k, bold=True)
        set_cell(c1, v)
    doc.add_paragraph()

    heading(doc, "1. Основание выполнения работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий отчет составлен в подтверждение выполнения работ по Этапу №3 "
        "«Подготовка к ИИ и системной модернизации» в соответствии с Техническим заданием "
        "(Приложение №1 к Договору).")

    heading(doc, "2. Цель выполнения этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Целью выполнения Этапа №3 являлось:", italic=True)
    for g in [
        "формирование структуры базы знаний;",
        "реализация поискового механизма по базе знаний;",
        "подготовка API-архитектуры для будущего подключения ИИ;",
        "настройка системы логирования;",
        "подготовка технической документации для дальнейшего этапа модернизации.",
    ]:
        bullet(doc, g)

    heading(doc, "3. Выполненные работы", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    heading(doc, "3.1. Формирование базы знаний", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Реализована структура хранения и управления базой знаний:", italic=True)
    bullet(doc, "таблица knowledge_items: id, title, slug, content (JSONB), category_id, tags[], department_id, author_id, status, published_at, metadata (JSONB);")
    bullet(doc, "таблица kb_categories: id, name, slug, description, parent_id, sort_order — иерархическая категоризация;")
    bullet(doc, "публичная страница /knowledge-base с боковым фильтром по категориям;")
    bullet(doc, "детальные статьи /knowledge-base/[slug] с prose-рендером, тегами, блоком «Читайте также»;")
    bullet(doc, "CMS-интерфейс управления базой знаний (/admin/knowledge) с WYSIWYG-редактором.")

    heading(doc, "3.2. Поисковый модуль", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Реализован полнотекстовый поисковый механизм:", italic=True)
    bullet(doc, "интеграция Meilisearch — высокопроизводительный поисковый движок с поддержкой русского языка;")
    bullet(doc, "GET /api/search?q= — внутренний поиск по новостям, материалам базы знаний и страницам;")
    bullet(doc, "GET /api/v1/search — публичный API поиска с фильтрацией по типу контента (type: news | knowledge | pages) и пагинацией;")
    bullet(doc, "компонент NewsSearch на странице новостей с мгновенным откликом (useTransition);")
    bullet(doc, "скрипт переиндексации pnpm search:reindex для обновления поискового индекса.")

    heading(doc, "3.3. API-архитектура для ИИ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Реализован публичный REST API /api/v1/ для интеграции внешних систем и ИИ:", italic=True)

    api_tbl = doc.add_table(rows=6, cols=3)
    api_tbl.style = 'Table Grid'
    for j, h in enumerate(["Эндпоинт", "Метод", "Описание"]):
        p = api_tbl.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(10); run.font.name = 'Times New Roman'
    api_data = [
        ("/api/v1/stats",        "GET",  "Агрегированная статистика центра (публикации, проекты, сотрудники)"),
        ("/api/v1/departments",  "GET",  "Список подразделений с руководителями"),
        ("/api/v1/knowledge",    "GET",  "Материалы базы знаний (q, category, department, limit, offset)"),
        ("/api/v1/search",       "GET",  "Полнотекстовый поиск по всему контенту портала"),
        ("/api/v1/ai/query",     "POST", "ИИ-ответчик: подготовленный эндпоинт с RAG-архитектурой"),
    ]
    for i, row_data in enumerate(api_data):
        for j, txt in enumerate(row_data):
            p = api_tbl.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    doc.add_paragraph()
    bullet(doc, "Аутентификация: Bearer-токен через заголовок Authorization;")
    bullet(doc, "CORS-заголовки для безопасного cross-origin доступа;")
    bullet(doc, "ISR-кэширование ответов (300 сек. по умолчанию);")
    bullet(doc, "GET /api/health — healthcheck с пингом БД и временем отклика.")

    heading(doc, "3.4. Система логирования", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    bullet(doc, "Таблица api_logs: timestamp, method, path, query_params, status_code, response_time_ms, ip_address, user_agent — логирование всех обращений к /api/v1/;")
    bullet(doc, "Таблица crosspost_log: content_type, content_id, platform, status, external_post_id, external_url, error_message, sent_at — журнал кросс-постинга.")

    heading(doc, "3.5. Техническая документация", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Подготовлена и передана Заказчику техническая документация в составе 6 документов:", italic=True)
    docs_tech = [
        ("docs/1-architecture.md",   "Архитектура системы — стек, маршруты, БД, схема взаимодействия компонентов, стратегии кэширования"),
        ("docs/2-api-reference.md",  "Справочник API — все эндпоинты /api/v1/ и внутренние API, параметры, примеры запросов и ответов"),
        ("docs/3-admin-guide.md",    "Руководство администратора CMS — работа с контентом, публикация, кросс-постинг, файловый менеджер"),
        ("docs/4-deploy-guide.md",   "Руководство по деплою — пошаговое развертывание на VPS Ubuntu, настройка Nginx, PM2, SSL, резервное копирование"),
        ("docs/5-knowledge-base.md", "Управление базой знаний — структура категорий, добавление материалов, поиск"),
        ("docs/6-ai-roadmap.md",     "Дорожная карта ИИ — план внедрения RAG, pgvector, Claude API; технические требования второго этапа"),
    ]
    dt = doc.add_table(rows=len(docs_tech)+1, cols=2)
    dt.style = 'Table Grid'
    for j, h in enumerate(["Файл", "Содержание"]):
        p = dt.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(10); run.font.name = 'Times New Roman'
    for i, (fname, desc) in enumerate(docs_tech):
        for j, txt in enumerate([fname, desc]):
            p = dt.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    doc.add_paragraph()

    heading(doc, "4. Достижение KPI", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    kpi3_tbl = doc.add_table(rows=5, cols=3)
    kpi3_tbl.style = 'Table Grid'
    for j, h in enumerate(["KPI", "Плановый показатель", "Факт / Статус"]):
        p = kpi3_tbl.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(10); run.font.name = 'Times New Roman'
    kpi3_data = [
        ("База знаний структурирована", "100%",
         "Реализована полная структура (knowledge_items + kb_categories) — Выполнено ✓"),
        ("Поиск работает корректно", "Корректная обработка запросов",
         "Meilisearch + /api/v1/search — поиск по всем типам контента — Выполнено ✓"),
        ("Документация передана Заказчику", "Документация передана",
         "6 технических документов в папке docs/ — Выполнено ✓"),
        ("Отсутствие критических ошибок", "0 критических ошибок",
         "0 ошибок, TypeScript проверка пройдена — Выполнено ✓"),
    ]
    for i, row_data in enumerate(kpi3_data):
        for j, txt in enumerate(row_data):
            p = kpi3_tbl.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    doc.add_paragraph()

    heading(doc, "5. Результаты этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "По итогам выполнения Этапа №3 Заказчику переданы:")
    bullet(doc, "Приложение № 1 — Описание поискового модуля и базы знаний;")
    bullet(doc, "Приложение № 2 — Описание API-архитектуры;")
    bullet(doc, "Приложение № 3 — Состав технической документации.")

    heading(doc, "6. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы по Этапу №3 «Подготовка к ИИ и системной модернизации» выполнены "
        "в полном объеме в соответствии с Договором №27/02/26-1 и Техническим заданием. "
        "Все KPI достигнуты. Система готова к подключению интеллектуального модуля "
        "на следующем этапе развития.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "Генеральный директор")
    normal(doc, "ООО «АСТРА ГРУПП»")
    doc.add_paragraph()
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 1 к Отчету № 3 ───────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 1 к Отчету № 3", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 3 от 09.03.2026 г.", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    heading(doc, "ОПИСАНИЕ ПОИСКОВОГО МОДУЛЯ И БАЗЫ ЗНАНИЙ")

    heading(doc, "1. База знаний", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Структура хранения данных:")
    bullet(doc, "Модель knowledge_items хранит материалы с заголовком, slug, JSONB-контентом (Tiptap-формат), категорией, тегами, статусом публикации и метаданными;")
    bullet(doc, "Модель kb_categories обеспечивает иерархическую структуру категорий с возможностью вложенности;")
    bullet(doc, "Поддерживается связь материалов с подразделениями и авторами;")
    bullet(doc, "Доступна фильтрация по категориям на публичной странице /knowledge-base.")

    heading(doc, "2. Поисковый модуль", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    bullet(doc, "Meilisearch — высокопроизводительный поисковый движок с нативной поддержкой русского языка и опечаток;")
    bullet(doc, "Индексация охватывает: новости (news), материалы базы знаний (knowledge_items), статические страницы (pages);")
    bullet(doc, "Поиск возвращает релевантные результаты с выделением совпадений;")
    bullet(doc, "Интерфейсный компонент NewsSearch реализует мгновенный отклик без перезагрузки страницы;")
    bullet(doc, "Скрипт pnpm search:reindex обеспечивает повторную индексацию при обновлении данных.")

    heading(doc, "3. Принцип работы поиска", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Пользователь вводит запрос → компонент обращается к /api/v1/search → "
        "сервер запрашивает Meilisearch → возвращает ранжированные результаты с "
        "учетом релевантности и опечаток → результаты отображаются мгновенно.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 2 к Отчету № 3 ───────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 2 к Отчету № 3", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 3 от 09.03.2026 г.", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    heading(doc, "ОПИСАНИЕ API-АРХИТЕКТУРЫ")

    heading(doc, "1. Общая концепция", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Публичный API /api/v1/ представляет собой RESTful-интерфейс, предназначенный для "
        "интеграции внешних систем, в том числе систем искусственного интеллекта, "
        "с информационной системой ГБУ «АНИЦ».")

    heading(doc, "2. Аутентификация", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    bullet(doc, "Bearer-токен в заголовке Authorization: Bearer <token>;")
    bullet(doc, "Токен задается в переменных окружения (.env.local → API_KEY);")
    bullet(doc, "Для запросов с корректным токеном устанавливаются CORS-заголовки;")
    bullet(doc, "Неавторизованные запросы получают ответ 401 Unauthorized.")

    heading(doc, "3. Эндпоинт ИИ-ответчика", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "POST /api/v1/ai/query")
    normal(doc, "Тело запроса:", italic=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('{ "query": "Какие арктические исследования ведет АНИЦ?", "limit": 5 }')
    run.font.size = Pt(10); run.font.name = 'Courier New'
    normal(doc, "Архитектура: запрос → поиск в базе знаний → передача контекста языковой модели → ответ.")
    normal(doc,
        "Эндпоинт реализован в виде инфраструктурной заготовки. Подключение языковой модели "
        "(Claude API / GPT) и RAG-алгоритм предусмотрены в следующем этапе развития системы "
        "согласно docs/6-ai-roadmap.md.")

    heading(doc, "4. Логирование", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Каждый запрос к /api/v1/ фиксируется в таблице api_logs с полями: timestamp, "
        "method, path, query_params, status_code, response_time_ms, ip_address, user_agent. "
        "Это обеспечивает полный аудит использования API и основу для аналитики запросов ИИ.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 3 к Отчету № 3 ───────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 3 к Отчету № 3", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 3 от 09.03.2026 г.", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()

    heading(doc, "СОСТАВ ТЕХНИЧЕСКОЙ ДОКУМЕНТАЦИИ")

    heading(doc, "1. Перечень документов", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    docs_full = [
        ("1", "docs/1-architecture.md", "Архитектура системы",
         "Описание технического стека, архитектуры Next.js App Router, структуры базы данных (13 таблиц), "
         "API-маршрутов, стратегий ISR-кэширования, схемы взаимодействия компонентов"),
        ("2", "docs/2-api-reference.md", "Справочник API",
         "Полное описание всех API-эндпоинтов: публичные /api/v1/ и внутренние /api/; "
         "методы, параметры, форматы запросов и ответов, примеры вызовов"),
        ("3", "docs/3-admin-guide.md", "Руководство администратора CMS",
         "Пошаговое руководство по работе с CMS-панелью: публикация новостей, "
         "управление базой знаний, кросс-постинг, загрузка файлов, настройки системы"),
        ("4", "docs/4-deploy-guide.md", "Руководство по деплою",
         "Инструкции по развертыванию системы на VPS Ubuntu: установка зависимостей, "
         "настройка PostgreSQL, Meilisearch, Nginx (SSL), PM2; скрипты резервного копирования"),
        ("5", "docs/5-knowledge-base.md", "Управление базой знаний",
         "Структура категорий базы знаний, добавление и редактирование материалов, "
         "теги, поиск, связи с подразделениями"),
        ("6", "docs/6-ai-roadmap.md", "Дорожная карта ИИ",
         "План внедрения интеллектуального модуля: RAG-архитектура, pgvector, "
         "интеграция Claude API, технические требования, оценка трудозатрат"),
    ]

    dtbl = doc.add_table(rows=len(docs_full)+1, cols=4)
    dtbl.style = 'Table Grid'
    for j, h in enumerate(["№", "Файл", "Название", "Содержание"]):
        p = dtbl.cell(0, j).paragraphs[0]
        run = p.add_run(h); run.bold = True
        run.font.size = Pt(10); run.font.name = 'Times New Roman'
    for i, row_data in enumerate(docs_full):
        for j, txt in enumerate(row_data):
            p = dtbl.cell(i+1, j).paragraphs[0]
            run = p.add_run(txt)
            run.font.size = Pt(9); run.font.name = 'Times New Roman'

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── СОСТАВ ДОКУМЕНТОВ ──────────────────────────────────────────────────
    header_line(doc, STAGE)
    doc.add_paragraph()
    heading(doc, "СОСТАВ ПЕРЕДАВАЕМЫХ ДОКУМЕНТОВ")
    heading(doc, "Этап № 3 – «Подготовка к ИИ и системной модернизации»")
    doc.add_paragraph()

    docs_list3 = [
        ("1)", "Акт № 3 сдачи-приемки выполненных работ к Договору № 27/02/26-1 от 27.02.2026 г. (является итоговым Актом по Договору)"),
        ("2)", "Отчет № 3 от 09.03.2026 г. о выполнении работ по Этапу № 3 «Подготовка к ИИ и системной модернизации»"),
        ("3)", "Приложение № 1 к Отчету № 3 по Этапу № 3 от 09.03.2026 г. – Описание поискового модуля и базы знаний"),
        ("4)", "Приложение № 2 к Отчету № 3 по Этапу № 3 от 09.03.2026 г. – Описание API-архитектуры"),
        ("5)", "Приложение № 3 к Отчету № 3 по Этапу № 3 от 09.03.2026 г. – Состав технической документации"),
    ]
    for num, text in docs_list3:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run_n = p.add_run(f"{num}\t")
        run_n.font.size = Pt(12); run_n.font.name = 'Times New Roman'
        run_t = p.add_run(text)
        run_t.font.size = Pt(12); run_t.font.name = 'Times New Roman'

    out = "/home/ash/Изображения/projects/gbu-anic-web-portal/docs/closing-docs/Комплект_документов_Этап3.docx"
    doc.save(out)
    print(f"Saved: {out}")
    return out


def build_stage1():
    doc = Document()
    set_margins(doc)
    STAGE = "ЭТАП № 1 – АУДИТ И ПРОЕКТИРОВАНИЕ"
    DATE  = "«01» марта 2026 г."

    # ── АКТ № 1 ────────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "АКТ № 1")
    heading(doc, "сдачи-приемки выполненных работ")
    heading(doc, "к Договору № 27/02/26-1 от 27.02.2026 г.")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    doc.add_paragraph()

    normal(doc,
        "Государственное бюджетное учреждение «Арктический научно-исследовательский "
        "центр Республики Саха (Якутия)», именуемое в дальнейшем «Заказчик», в лице "
        "и.о. директора Хохолова Артура Аркадьевича, действующего на основании приказа "
        "Минобрнауки РС(Я) от 15.01.2026г. №09-22/9, доверенности № 23 от 30.01.2026 г., "
        "с одной стороны, и")
    normal(doc,
        "Общество с ограниченной ответственностью «АСТРА ГРУПП», именуемое в дальнейшем "
        "«Исполнитель», в лице генерального директора Ядрихинского Сергея Александровича, "
        "действующего на основании Устава, с другой стороны,")
    normal(doc, "составили настоящий Акт о нижеследующем:", italic=True)

    heading(doc, "1. Основание", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий Акт составлен во исполнение Договора № 27/02/26-1 от 27.02.2026 г. и "
        "Технического задания (Приложение № 1 к Договору).")

    heading(doc, "2. Перечень выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Исполнителем выполнены работы по Этапу №1 «Аудит и проектирование» в соответствии "
        "с Техническим заданием (Приложение №1 к Договору), а именно:")

    works1 = [
        "проведен анализ текущей архитектуры информационной системы Заказчика;",
        "проведен анализ структуры сайта и контента;",
        "выявлены технические и логические ограничения действующей системы;",
        "проведен базовый аудит производительности;",
        "разработана архитектурная схема модернизации;",
        "разработана обновленная структура сайта (Sitemap);",
        "разработана концепция развития системы управления контентом (CMS);",
        "разработана концепция внедрения интеллектуального модуля.",
    ]
    for w in works1:
        bullet(doc, w)

    normal(doc, "\nРезультаты работ переданы Заказчику в составе комплекта документов.")

    heading(doc, "3. Стоимость выполненных работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Стоимость работ по Этапу №1 составляет: 45 000 (сорок пять тысяч) рублей 00 копеек. "
        "НДС не облагается в связи с применением Исполнителем АУСН.")

    heading(doc, "4. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы выполнены в полном объеме, в соответствии с условиями Договора и Технического "
        "задания. Работы выполнены в соответствии с Календарным планом.")
    normal(doc, "Заказчик претензий к объему, качеству и срокам выполнения работ не имеет.")
    normal(doc, "Настоящий Акт является основанием для оплаты выполненных работ.")

    heading(doc, "5. Подписи сторон", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    doc.add_paragraph()
    signature_table(doc)

    pagebreak(doc)

    # ── ОТЧЕТ № 1 ──────────────────────────────────────────────────────────
    header_line(doc, STAGE)
    heading(doc, "ОТЧЕТ № 1")
    heading(doc, "о выполнении работ по Этапу № 1")
    heading(doc, "«Аудит и проектирование»")
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("г. Якутск")
    run.font.size = Pt(12); run.font.name = 'Times New Roman'
    tab = p.add_run("\t\t\t\t\t\t\t\t")
    tab.font.size = Pt(12)
    run2 = p.add_run(DATE)
    run2.font.size = Pt(12); run2.font.name = 'Times New Roman'

    doc.add_paragraph()
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = 'Table Grid'
    rows_data = [
        ("Заказчик:",            "ГБУ «Арктический научно-исследовательский центр Республики Саха (Якутия)»"),
        ("Исполнитель:",         "ООО «АСТРА ГРУПП»"),
        ("Основание:",           "Договор № 27/02/26-1 от 27.02.2026 г."),
        ("Период выполнения:",   "27.02.2026 г. – 01.03.2026 г."),
    ]
    for i, (k, v) in enumerate(rows_data):
        c0 = tbl.cell(i, 0); c1 = tbl.cell(i, 1)
        def set_cell(cell, text, bold=False):
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.bold = bold; run.font.size = Pt(11); run.font.name = 'Times New Roman'
        set_cell(c0, k, bold=True)
        set_cell(c1, v)
    doc.add_paragraph()

    heading(doc, "1. Основание выполнения работ", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Настоящий отчет составлен в подтверждение выполнения работ по Этапу №1 "
        "«Аудит и проектирование» в соответствии с Техническим заданием (Приложение №1 к Договору).")
    normal(doc,
        "Этап №1 предусматривает проведение комплексного анализа действующей информационной "
        "системы Заказчика и разработку концепции ее модернизации.")

    heading(doc, "2. Цель выполнения этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Целью выполнения Этапа №1 являлось:", italic=True)
    for g in [
        "проведение комплексного анализа действующей информационной системы Заказчика;",
        "выявление технических и логических ограничений;",
        "формирование архитектурной модели модернизации;",
        "разработка обновленной структуры сайта;",
        "подготовка концепции дальнейшего развития системы управления контентом (CMS);",
        "подготовка концепции интеграции интеллектуального модуля.",
    ]:
        bullet(doc, g)

    heading(doc, "3. Объект аудита", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Объектом аудита являлась действующая информационная система Заказчика, включая:", italic=True)
    for item in [
        "публичную часть портала (домен аниц.рф);",
        "структуру навигации;",
        "административную часть (при наличии доступа);",
        "модель хранения и публикации контента;",
        "логическую структуру разделов;",
        "производительность и стабильность загрузки страниц.",
    ]:
        bullet(doc, item)

    heading(doc, "4. Выполненные работы", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    heading(doc, "4.1. Анализ текущей архитектуры Системы", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "В рамках анализа архитектуры выполнено:", italic=True)
    for item in [
        "изучение общей структуры портала;",
        "анализ логики маршрутизации;",
        "анализ структуры представления данных;",
        "оценка масштабируемости;",
        "оценка возможности внедрения дополнительных модулей.",
    ]:
        bullet(doc, item)
    normal(doc,
        "Вывод: текущая архитектура требует систематизации и перехода к модульной структуре "
        "с четким разделением публичной и административной части.", italic=True)

    heading(doc, "4.2. Анализ структуры сайта и контента", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Проведен анализ:", italic=True)
    for item in ["иерархии разделов;", "логики распределения информации;",
                 "полноты и доступности ключевых разделов;", "пользовательского пути (user flow)."]:
        bullet(doc, item)
    normal(doc, "Выявлены следующие направления улучшения:", italic=True)
    for item in ["необходимость оптимизации навигации;",
                 "структурирование информационных блоков;",
                 "подготовка базы для расширения разделов."]:
        bullet(doc, item)

    heading(doc, "4.3. Выявление технических и логических ограничений", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Определены:", italic=True)
    for item in [
        "ограничения текущей модели масштабирования;",
        "недостаточная модульность системы;",
        "отсутствие подготовленной архитектуры для интеграции интеллектуальных сервисов;",
        "необходимость оптимизации производительности.",
    ]:
        bullet(doc, item)

    heading(doc, "4.4. Базовый аудит производительности", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Проведена оценка:", italic=True)
    for item in ["времени загрузки основных страниц;",
                 "корректности загрузки статических и динамических ресурсов;",
                 "общей устойчивости работы портала."]:
        bullet(doc, item)
    normal(doc, "Сформированы рекомендации по:", italic=True)
    for item in ["внедрению механизмов кэширования;",
                 "оптимизации рендеринга;",
                 "сокращению избыточных запросов."]:
        bullet(doc, item)

    heading(doc, "4.5. Разработка архитектурной схемы модернизации", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Разработана целевая архитектурная модель, включающая:", italic=True)
    for item in [
        "разделение публичной и административной частей;",
        "централизованную систему управления контентом (CMS);",
        "API-архитектуру;",
        "подготовку к интеграции интеллектуального модуля;",
        "возможность горизонтального масштабирования.",
    ]:
        bullet(doc, item)
    normal(doc, "Архитектурная схема разработана и представлена Заказчику для согласования.", italic=True)

    heading(doc, "4.6. Разработка обновленной структуры сайта (Sitemap)", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Сформирована обновленная структура публичной части портала, включающая:", italic=True)
    for item in ["логически сгруппированные разделы;", "понятную иерархию;",
                 "упрощенную навигацию;", "основу для дальнейшего расширения."]:
        bullet(doc, item)
    normal(doc, "Разработана схема структуры сайта (Sitemap) в текстовом формате.", italic=True)

    heading(doc, "4.7. Подготовка концепции развития CMS", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Разработана концепция:", italic=True)
    for item in [
        "централизованного управления контентом;",
        "разграничения ролей пользователей;",
        "управления новостями, публикациями, документами;",
        "последующего масштабирования системы.",
    ]:
        bullet(doc, item)

    heading(doc, "4.8. Подготовка концепции внедрения интеллектуального модуля", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Разработана концепция архитектурной интеграции интеллектуального модуля, предусматривающая:", italic=True)
    for item in [
        "формирование базы знаний;",
        "создание API-слоя;",
        "возможность подключения поискового механизма;",
        "возможность последующей интеграции языковых моделей.",
    ]:
        bullet(doc, item)
    normal(doc, "Концепция носит архитектурный характер и соответствует требованиям ТЗ.", italic=True)

    heading(doc, "5. Результаты этапа", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "По итогам выполнения Этапа №1 Заказчику переданы:")
    for item in [
        "Приложение № 1 — Отчет об аудите действующей Системы;",
        "Приложение № 2 — Архитектурная схема модернизации;",
        "Приложение № 3 — Обновленная структура сайта (Sitemap);",
        "Приложение № 4 — Концепция дальнейшего развития Системы.",
    ]:
        bullet(doc, item)

    heading(doc, "6. Заключение", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Работы по Этапу №1 «Аудит и проектирование» выполнены в полном объеме в соответствии "
        "с Договором №27/02/26-1 и Техническим заданием.")

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "Генеральный директор")
    normal(doc, "ООО «АСТРА ГРУПП»")
    doc.add_paragraph()
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 1 ─────────────────────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 1 к Отчету № 1", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 1 от 01.03.2026 г.",  italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    heading(doc, "ОТЧЕТ ОБ АУДИТЕ ДЕЙСТВУЮЩЕЙ ИНФОРМАЦИОННОЙ СИСТЕМЫ")
    heading(doc, "по состоянию на дату проведения работ", size=12, bold=False)

    heading(doc, "1. Общие сведения", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Объект аудита: информационная система ГБУ «АНИЦ» (публичный портал www.аниц.рф). "
        "Цель аудита: оценка текущего состояния архитектуры, структуры, производительности "
        "и масштабируемости системы.")

    heading(doc, "2. Анализ архитектуры", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "В рамках аудита выполнено:", italic=True)
    for item in ["анализ общей логики построения портала;",
                 "анализ структуры маршрутов;",
                 "оценка модульности архитектуры;",
                 "оценка возможности интеграции дополнительных сервисов."]:
        bullet(doc, item)
    normal(doc, "Выводы:", italic=True)
    for item in [
        "архитектура носит линейный характер;",
        "отсутствует четкое разделение уровней (публичная часть / API / администрирование);",
        "масштабируемость ограничена текущей структурой;",
        "архитектура требует перехода к модульному принципу.",
    ]:
        bullet(doc, item)

    heading(doc, "3. Анализ структуры сайта и контента", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Проведен анализ:", italic=True)
    for item in ["структуры разделов;", "навигационной логики;",
                 "иерархии страниц;", "пользовательского пути."]:
        bullet(doc, item)
    normal(doc, "Выявленные особенности:", italic=True)
    for item in [
        "избыточная вложенность отдельных разделов;",
        "необходимость упорядочивания контента;",
        "отсутствие системной базы знаний;",
        "необходимость унификации структуры страниц.",
    ]:
        bullet(doc, item)

    heading(doc, "4. Анализ производительности", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Проведена базовая оценка:", italic=True)
    for item in ["времени загрузки страниц;",
                 "стабильности отображения;",
                 "корректности загрузки ресурсов."]:
        bullet(doc, item)
    normal(doc, "Выводы:", italic=True)
    for item in ["возможна оптимизация скорости загрузки;",
                 "требуется внедрение механизмов кэширования;",
                 "требуется оптимизация структуры отдачи контента."]:
        bullet(doc, item)

    heading(doc, "5. Выявленные ограничения", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in [
        "отсутствие подготовленной архитектуры для ИИ-интеграции;",
        "отсутствие централизованной модели управления контентом;",
        "ограниченные возможности масштабирования;",
        "необходимость модернизации структуры CMS.",
    ]:
        bullet(doc, item)

    heading(doc, "6. Итог аудита", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Действующая информационная система функционирует стабильно, однако требует "
        "архитектурной модернизации для:", italic=True)
    for item in ["повышения масштабируемости;", "внедрения интеллектуального модуля;",
                 "оптимизации производительности;", "структурирования контента."]:
        bullet(doc, item)

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 2 ─────────────────────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 2 к Отчету № 1", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 1 от 01.03.2026 г.",  italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    heading(doc, "АРХИТЕКТУРНАЯ СХЕМА МОДЕРНИЗАЦИИ")

    heading(doc, "1. Общая концепция", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Предлагаемая архитектура основана на модульном принципе и включает следующие уровни:", italic=True)
    for i, item in enumerate([
        "Публичный уровень (Frontend);",
        "Уровень API;",
        "Административный уровень (CMS);",
        "Уровень хранения данных (База данных);",
        "Уровень поискового и интеллектуального модуля.",
    ], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}. {item}")
        run.font.size = Pt(12); run.font.name = 'Times New Roman'

    heading(doc, "2. Логическая модель взаимодействия", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    schema_tbl = doc.add_table(rows=1, cols=1)
    schema_tbl.style = 'Table Grid'
    cell = schema_tbl.cell(0, 0)
    for line in [
        "Пользователь",
        "↓",
        "Публичный портал",
        "↓",
        "API-слой",
        "↓",
        "База данных / Поисковый модуль",
        "↓",
        "Административная панель",
    ]:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(12); run.font.name = 'Times New Roman'
    doc.add_paragraph()

    heading(doc, "3. Принципы модернизации", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in [
        "разделение ответственности модулей;",
        "масштабируемость;",
        "возможность горизонтального расширения;",
        "подготовка к интеграции интеллектуального модуля;",
        "централизованное управление контентом.",
    ]:
        bullet(doc, item)

    heading(doc, "4. Результат архитектурного проектирования", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Разработанная схема обеспечивает:")
    for item in [
        "возможность интеграции поискового модуля;",
        "возможность внедрения ИИ-сервиса;",
        "централизованную систему управления контентом;",
        "гибкость масштабирования.",
    ]:
        bullet(doc, item)

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 3 ─────────────────────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 3 к Отчету № 1", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 1 от 01.03.2026 г.",  italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    heading(doc, "ОБНОВЛЕННАЯ СТРУКТУРА САЙТА (SITEMAP)")

    heading(doc, "1. Основные разделы публичного портала", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in [
        "Главная страница",
        "О центре",
        "Научная деятельность",
        "Новости",
        "Медиа",
        "Образование",
        "Партнёры",
        "Документы",
        "Закупки",
        "Контакты",
    ]:
        bullet(doc, item)
    normal(doc, "Дополнительно:", italic=True)
    bullet(doc, "База знаний (раздел для структурированной информации)")

    heading(doc, "2. Принципы построения структуры", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in [
        "логическая группировка разделов;",
        "минимизация избыточной вложенности;",
        "доступность ключевых разделов в 1–2 клика;",
        "подготовка к расширению структуры.",
    ]:
        bullet(doc, item)

    heading(doc, "3. Административная структура (CMS)", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc, "Предусмотрены разделы управления:", italic=True)
    for item in ["Новости;", "Документы;", "Партнёры;", "Медиа;", "База знаний;", "Пользователи и роли."]:
        bullet(doc, item)

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── Приложение № 4 ─────────────────────────────────────────────────────
    header_line(doc, STAGE)
    normal(doc, "Приложение № 4 к Отчету № 1", italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    normal(doc, "к Этапу № 1 от 01.03.2026 г.",  italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_paragraph()
    heading(doc, "КОНЦЕПЦИЯ ДАЛЬНЕЙШЕГО РАЗВИТИЯ СИСТЕМЫ")

    heading(doc, "1. Цель развития", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    normal(doc,
        "Создание масштабируемой, устойчивой и расширяемой цифровой платформы научного центра.")

    heading(doc, "2. Направления развития", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)

    heading(doc, "2.1. Развитие архитектуры", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in ["внедрение поискового сервиса;",
                 "оптимизация производительности;",
                 "развитие API-инфраструктуры."]:
        bullet(doc, item)

    heading(doc, "2.2. Развитие CMS", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in ["расширение функционала администрирования;",
                 "внедрение аналитики;",
                 "автоматизация публикации."]:
        bullet(doc, item)

    heading(doc, "2.3. Интеллектуальный модуль", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in ["формирование базы знаний;",
                 "интеграция поискового механизма;",
                 "внедрение ИИ-помощника;",
                 "развитие RAG-архитектуры."]:
        bullet(doc, item)

    heading(doc, "2.4. Дизайн и пользовательский опыт", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in ["разработка дизайн-концепции;",
                 "унификация визуального стиля;",
                 "повышение доступности."]:
        bullet(doc, item)

    heading(doc, "3. Ожидаемый эффект", align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True)
    for item in [
        "повышение удобства для пользователей;",
        "повышение управляемости контента;",
        "подготовка к цифровой трансформации;",
        "соответствие современным требованиям веб-инфраструктуры.",
    ]:
        bullet(doc, item)

    doc.add_paragraph()
    normal(doc, "Исполнитель:", bold=True)
    normal(doc, "_____________ / Ядрихинский С.А.")
    normal(doc, "М.П.")

    pagebreak(doc)

    # ── СОСТАВ ДОКУМЕНТОВ ──────────────────────────────────────────────────
    header_line(doc, STAGE)
    doc.add_paragraph()
    heading(doc, "СОСТАВ ПЕРЕДАВАЕМЫХ ДОКУМЕНТОВ")
    heading(doc, "Этап № 1 – «Аудит и проектирование»")
    doc.add_paragraph()

    docs_list1 = [
        ("1)", "Акт № 1 сдачи-приемки выполненных работ к Договору № 27/02/26-1 от 27.02.2026 г."),
        ("2)", "Отчет № 1 от 01.03.2026 г. о выполнении работ по Этапу № 1 «Аудит и проектирование»"),
        ("3)", "Приложение № 1 к Отчету № 1 по Этапу № 1 от 01.03.2026 г. – Отчет об аудите действующей информационной системы"),
        ("4)", "Приложение № 2 к Отчету № 1 по Этапу № 1 от 01.03.2026 г. – Архитектурная схема модернизации"),
        ("5)", "Приложение № 3 к Отчету № 1 по Этапу № 1 от 01.03.2026 г. – Обновленная структура сайта (Sitemap)"),
        ("6)", "Приложение № 4 к Отчету № 1 по Этапу № 1 от 01.03.2026 г. – Концепция дальнейшего развития системы"),
    ]
    for num, text in docs_list1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        run_n = p.add_run(f"{num}\t")
        run_n.font.size = Pt(12); run_n.font.name = 'Times New Roman'
        run_t = p.add_run(text)
        run_t.font.size = Pt(12); run_t.font.name = 'Times New Roman'

    out = "/home/ash/Изображения/projects/gbu-anic-web-portal/docs/closing-docs/Комплект_документов_Этап1.docx"
    doc.save(out)
    print(f"Saved: {out}")
    return out


if __name__ == "__main__":
    build_stage1()
    build_stage2()
    build_stage3()
    print("Done!")
